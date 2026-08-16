import os
import json
import uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify
from dotenv import load_dotenv
import requests

load_dotenv()

app = Flask(__name__)

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
FILES_DIR = BASE_DIR / "files"

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"


def load_json(path, default=None):
    if default is None:
        default = {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default


def save_json(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_system_prompt():
    data = load_json(DATA_DIR / "system_prompt.json",
                     {"role": "system", "content": "You are Phantom Intelligence."})
    return {"role": "system", "content": data.get("content", "You are Phantom Intelligence, a helpful AI assistant.")}


def load_db():
    return load_json(DATA_DIR / "history.json", {"conversations": [], "active": None})


def save_db(db):
    save_json(DATA_DIR / "history.json", db)


def get_active_conversation(db):
    """Return the active conversation dict, or create a new one if none active or hidden."""
    active_id = db.get("active")
    if active_id:
        for conv in db["conversations"]:
            if conv["id"] == active_id and not conv.get("hidden"):
                return conv
    # No active visible conversation — create one
    conv = {
        "id": str(uuid.uuid4()),
        "title": "New Chat",
        "messages": [],
        "created": datetime.now().isoformat(),
        "updated": datetime.now().isoformat()
    }
    db["conversations"].append(conv)
    db["active"] = conv["id"]
    save_db(db)
    return conv


# ─── Tool Definitions ───────────────────────────────────────────────────

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "Get the current date and time.",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_location",
            "description": "Get the user's approximate location based on IP address.",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_weather",
            "description": "Get accurate current weather for any city worldwide. Returns temperature, humidity, wind, condition, and a Windy map link.",
            "parameters": {
                "type": "object",
                "properties": {
                    "city": {"type": "string", "description": "City name, e.g. Beijing, London, Tokyo"}
                },
                "required": ["city"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read the content of a file from the files directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Name of the file to read"}
                },
                "required": ["filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Write content to a file in the files directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Name of the file to write"},
                    "content": {"type": "string", "description": "Content to write"}
                },
                "required": ["filename", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_files",
            "description": "List all files in the files directory.",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Search the web using DuckDuckGo.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Search query"},
                    "max_results": {"type": "integer", "description": "Max results", "default": 5}
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "download_file",
            "description": "Download a file from a URL to the files directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL to download"},
                    "filename": {"type": "string", "description": "Save as filename (optional)"}
                },
                "required": ["url"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "Generate an image from a text description. The image is created from scratch based on the prompt. Returns the path to the saved image file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Text description of the image to generate. Be detailed and specific. Include style, composition, colors, mood, lighting."},
                    "size": {"type": "string", "description": "Image size: square, square_hd, portrait_4_3, portrait_16_9, landscape_4_3, landscape_16_9. Default: square."}
                },
                "required": ["prompt"]
            }
        }
    }
]


# ─── Tool Executors ─────────────────────────────────────────────────────

def execute_tool(tool_name, arguments):
    try:
        if tool_name == "get_current_time":
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            return {"success": True, "result": f"Current time: {now}"}

        elif tool_name == "get_current_location":
            try:
                resp = requests.get("http://ip-api.com/json/", timeout=5)
                data = resp.json()
                if data.get("status") == "success":
                    return {"success": True, "result": json.dumps({
                        "city": data.get("city"),
                        "region": data.get("regionName"),
                        "country": data.get("country"),
                        "lat": data.get("lat"),
                        "lon": data.get("lon"),
                        "timezone": data.get("timezone")
                    }, ensure_ascii=False)}
                return {"success": False, "result": "Unable to determine location."}
            except Exception as e:
                return {"success": False, "result": f"Location lookup failed: {str(e)}"}

        elif tool_name == "get_weather":
            city = arguments.get("city", "Beijing")
            try:
                from urllib.parse import quote
                import time
                
                # Step 1: Geocode city → lat/lon using Open-Meteo Geocoding API (with retry)
                geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={quote(city)}&count=1&language=en&format=json"
                geo_data = None
                for attempt in range(2):
                    try:
                        geo_resp = requests.get(geo_url, timeout=15)
                        geo_data = geo_resp.json()
                        break
                    except requests.exceptions.Timeout:
                        if attempt == 0:
                            time.sleep(1)
                        continue
                
                if not geo_data or not geo_data.get("results"):
                    # Fallback to wttr.in if geocoding fails
                    try:
                        fallback = requests.get(f"https://wttr.in/{quote(city)}?format=%C+%t+%h+%w&m", timeout=10)
                        windy_url = f"https://www.windy.com/?search={quote(city)}"
                        return {"success": True, "result": f"Weather in {city}: {fallback.text.strip()}\nWindy map: {windy_url}"}
                    except:
                        return {"success": False, "result": f"Could not find weather for '{city}'. Check the city name or try again."}
                
                result = geo_data["results"][0]
                lat = result["latitude"]
                lon = result["longitude"]
                found_name = result.get("name", city)
                country = result.get("country", "")
                
                # Step 2: Get current weather from Open-Meteo (with retry)
                weather_url = (
                    f"https://api.open-meteo.com/v1/forecast"
                    f"?latitude={lat}&longitude={lon}"
                    f"&current=temperature_2m,relative_humidity_2m,apparent_temperature"
                    f",weather_code,wind_speed_10m,wind_direction_10m,is_day"
                    f"&timezone=auto"
                )
                w_data = None
                for attempt in range(2):
                    try:
                        w_resp = requests.get(weather_url, timeout=15)
                        w_data = w_resp.json()
                        break
                    except requests.exceptions.Timeout:
                        if attempt == 0:
                            time.sleep(1)
                        continue
                
                if not w_data:
                    # Fallback to wttr.in
                    try:
                        fallback = requests.get(f"https://wttr.in/{quote(city)}?format=%C+%t+%h+%w&m", timeout=10)
                        windy_url = f"https://www.windy.com/?{lat},{lon},8"
                        return {"success": True, "result": f"Weather in {found_name}: {fallback.text.strip()}\nWindy map: {windy_url}"}
                    except:
                        return {"success": False, "result": f"Weather service is slow. Open Windy directly: https://www.windy.com/?{lat},{lon},8"}
                
                current = w_data.get("current", {})
                
                # Weather code → description (WMO codes)
                wmo_codes = {
                    0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
                    45: "Foggy", 48: "Depositing rime fog",
                    51: "Light drizzle", 53: "Moderate drizzle", 55: "Dense drizzle",
                    56: "Light freezing drizzle", 57: "Dense freezing drizzle",
                    61: "Slight rain", 63: "Moderate rain", 65: "Heavy rain",
                    66: "Light freezing rain", 67: "Heavy freezing rain",
                    71: "Slight snow", 73: "Moderate snow", 75: "Heavy snow",
                    77: "Snow grains",
                    80: "Slight rain showers", 81: "Moderate rain showers", 82: "Violent rain showers",
                    85: "Slight snow showers", 86: "Heavy snow showers",
                    95: "Thunderstorm", 96: "Thunderstorm with slight hail", 99: "Thunderstorm with heavy hail"
                }
                weather_code = current.get("weather_code", 0)
                condition = wmo_codes.get(weather_code, f"Code {weather_code}")
                
                temp = current.get("temperature_2m", "N/A")
                feels = current.get("apparent_temperature", "N/A")
                humidity = current.get("relative_humidity_2m", "N/A")
                wind_speed = current.get("wind_speed_10m", "N/A")
                wind_dir = current.get("wind_direction_10m", "N/A")
                is_day = current.get("is_day", 1)
                
                # Wind direction → compass
                dirs = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
                        "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
                wind_compass = dirs[round(wind_dir / 22.5) % 16] if isinstance(wind_dir, (int, float)) else "N/A"
                
                windy_url = f"https://www.windy.com/?{lat},{lon},8"
                
                result_text = (
                    f"Weather in {found_name}, {country}:\n"
                    f"Condition: {condition}\n"
                    f"Temperature: {temp}°C (feels like {feels}°C)\n"
                    f"Humidity: {humidity}%\n"
                    f"Wind: {wind_speed} km/h from {wind_compass} ({wind_dir}°)\n"
                    f"Day/Night: {'Daytime' if is_day else 'Nighttime'}\n"
                    f"Windy map: {windy_url}"
                )
                return {"success": True, "result": result_text}
            except Exception as e:
                return {"success": False, "result": f"Weather lookup failed: {str(e)}"}

        elif tool_name == "read_file":
            filename = arguments.get("filename", "")
            filepath = (FILES_DIR / filename).resolve()
            if not str(filepath).startswith(str(FILES_DIR.resolve())):
                return {"success": False, "result": "Error: Invalid file path."}
            if not filepath.exists():
                return {"success": False, "result": f"Error: File '{filename}' not found."}
            content = filepath.read_text(encoding="utf-8", errors="replace")
            return {"success": True, "result": content}

        elif tool_name == "write_file":
            filename = arguments.get("filename", "")
            content = arguments.get("content", "")
            filepath = (FILES_DIR / filename).resolve()
            if not str(filepath).startswith(str(FILES_DIR.resolve())):
                return {"success": False, "result": "Error: Invalid file path."}
            filepath.parent.mkdir(parents=True, exist_ok=True)
            filepath.write_text(content, encoding="utf-8")
            return {"success": True, "result": f"File '{filename}' written ({len(content)} chars)."}

        elif tool_name == "list_files":
            if not FILES_DIR.exists():
                return {"success": True, "result": "No files found."}
            files = [str(f.relative_to(FILES_DIR)) for f in FILES_DIR.rglob("*") if f.is_file()]
            if not files:
                return {"success": True, "result": "No files found."}
            return {"success": True, "result": "Files:\n" + "\n".join(files)}

        elif tool_name == "web_search":
            query = arguments.get("query", "")
            max_results = arguments.get("max_results", 5)
            try:
                from ddgs import DDGS
                results = []
                with DDGS() as ddgs:
                    for r in ddgs.text(query, max_results=max_results):
                        results.append(f"{r.get('title','')}\n{r.get('href','')}\n{r.get('body','')}")
                if not results:
                    return {"success": True, "result": f"No results for '{query}'."}
                return {"success": True, "result": "\n\n".join(results)}
            except Exception as e:
                return {"success": False, "result": f"Search failed: {str(e)}"}

        elif tool_name == "download_file":
            url = arguments.get("url", "")
            filename = arguments.get("filename") or url.split("/")[-1].split("?")[0] or "download"
            filepath = (FILES_DIR / filename).resolve()
            if not str(filepath).startswith(str(FILES_DIR.resolve())):
                return {"success": False, "result": "Error: Invalid file path."}
            try:
                resp = requests.get(url, timeout=30, stream=True)
                resp.raise_for_status()
                filepath.parent.mkdir(parents=True, exist_ok=True)
                with open(filepath, "wb") as f:
                    for chunk in resp.iter_content(chunk_size=8192):
                        f.write(chunk)
                size = filepath.stat().st_size
                return {"success": True, "result": f"Downloaded '{filename}' ({size} bytes)."}
            except Exception as e:
                return {"success": False, "result": f"Download failed: {str(e)}"}

        elif tool_name == "generate_image":
            prompt = arguments.get("prompt", "")
            size = arguments.get("size", "square")
            valid_sizes = {"square", "square_hd", "portrait_4_3", "portrait_16_9", "landscape_4_3", "landscape_16_9"}
            if size not in valid_sizes:
                size = "square"

            from urllib.parse import quote
            import time
            import hashlib

            # Map sizes to dimensions
            size_map = {
                "square": (1024, 1024),
                "square_hd": (1536, 1536),
                "portrait_4_3": (768, 1024),
                "portrait_16_9": (576, 1024),
                "landscape_4_3": (1024, 768),
                "landscape_16_9": (1024, 576),
            }
            w, h = size_map.get(size, (1024, 1024))
            
            # Use Pollinations.ai free API
            url = f"https://image.pollinations.ai/prompt/{quote(prompt)}?width={w}&height={h}&model=flux&nologo=true"
            try:
                resp = requests.get(url, timeout=120, allow_redirects=True)
                resp.raise_for_status()
                ct = resp.headers.get("Content-Type", "").lower()
                content = resp.content
                
                # Detect image type from magic bytes
                is_jpg = content[:4] == b'\xff\xd8\xff\xe0' or content[:4] == b'\xff\xd8\xff\xe1' or content[:3] == b'\xff\xd8\xff'
                is_png = content[:8] == b'\x89PNG\r\n\x1a\n'
                is_gif = content[:6] in (b'GIF89a', b'GIF87a')
                is_webp = content[:4] == b'RIFF' and content[8:12] == b'WEBP'
                is_image = "image" in ct or is_jpg or is_png or is_gif or is_webp
                
                if is_image and len(content) > 1000:
                    if is_png:
                        ext = ".png"
                        mime = "image/png"
                    elif is_gif:
                        ext = ".gif"
                        mime = "image/gif"
                    elif is_webp:
                        ext = ".webp"
                        mime = "image/webp"
                    else:
                        ext = ".jpg"
                        mime = "image/jpeg"
                    filehash = hashlib.md5(prompt.encode()).hexdigest()[:8]
                    filename = f"generated_{filehash}{ext}"
                    filepath = (FILES_DIR / filename).resolve()
                    filepath.parent.mkdir(parents=True, exist_ok=True)
                    with open(filepath, "wb") as f:
                        f.write(content)
                    size_bytes = filepath.stat().st_size
                    import base64
                    b64 = base64.b64encode(content).decode("utf-8")
                    data_url = f"data:{mime};base64,{b64}"
                    return {"success": True, "result": f"Image generated and saved as '{filename}' ({size_bytes} bytes). Prompt: {prompt}", "filename": filename, "data_url": data_url}
                else:
                    preview = content[:500].decode("utf-8", errors="replace")
                    return {"success": False, "result": f"Image API returned non-image content ({len(content)} bytes). Preview: {preview}"}
            except Exception as e:
                return {"success": False, "result": f"Image generation failed: {str(e)}"}

        else:
            return {"success": False, "result": f"Unknown tool: {tool_name}"}

    except Exception as e:
        return {"success": False, "result": f"Tool error: {str(e)}"}


# ─── DeepSeek Call ──────────────────────────────────────────────────────

def call_deepseek(messages, tools=None, tool_choice=None):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {"model": "deepseek-chat", "messages": messages, "stream": False}
    if tools:
        payload["tools"] = tools
    if tool_choice:
        payload["tool_choice"] = tool_choice
    resp = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()


# ─── Routes ─────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/conversations/new", methods=["POST"])
def new_conversation():
    db = load_db()
    conv = {
        "id": str(uuid.uuid4()),
        "title": "New Chat",
        "messages": [],
        "created": datetime.now().isoformat(),
        "updated": datetime.now().isoformat()
    }
    db["conversations"].append(conv)
    db["active"] = conv["id"]
    save_db(db)
    return jsonify({"conversation_id": conv["id"]})


@app.route("/api/conversations/switch", methods=["POST"])
def switch_conversation():
    data = request.get_json(force=True)
    conv_id = data.get("conversation_id", "")
    if not conv_id:
        return jsonify({"error": "conversation_id required"}), 400

    db = load_db()
    for conv in db["conversations"]:
        if conv["id"] == conv_id:
            db["active"] = conv_id
            save_db(db)
            return jsonify({"status": "ok", "conversation": conv})

    return jsonify({"error": "Conversation not found"}), 404


@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True)
    user_message = data.get("message", "").strip()
    if not user_message:
        return jsonify({"error": "Message is required"}), 400

    db = load_db()
    conv = get_active_conversation(db)

    # Detect if user is asking for an image
    img_keywords = ["generate", "create", "draw", "make", "paint", "design", "生成", "画", "创建", "制作", "picture", "image", "photo", "illustration", "art", "图片", "照片", "图像", "绘画"]
    wants_image = any(kw in user_message.lower() for kw in img_keywords)

    # Build messages: system prompt + ALL conversations (hidden + visible) + current message
    messages = [load_system_prompt()]
    # Include ALL conversations for AI memory (including hidden ones)
    for c in db["conversations"]:
        for entry in c.get("messages", []):
            messages.append({"role": "user", "content": entry["user"]})
            messages.append({"role": "assistant", "content": entry["assistant"]})
    messages.append({"role": "user", "content": user_message})

    chat_request = {
        "model": "deepseek-chat",
        "messages": [dict(m) for m in messages],
        "tools": TOOLS
    }

    request_response_pairs = []
    image_tool_called = False

    max_loops = 30
    for loop_idx in range(max_loops):
        try:
            # Force generate_image tool on first call if user wants an image
            tc = None
            if loop_idx == 0 and wants_image:
                tc = {"type": "function", "function": {"name": "generate_image"}}
            result = call_deepseek(messages, TOOLS, tool_choice=tc)
        except requests.exceptions.HTTPError as e:
            error_body = ""
            try:
                error_body = e.response.text
            except Exception:
                pass
            return jsonify({
                "error": f"DeepSeek API error: {e.response.status_code}",
                "detail": error_body
            }), 502
        except Exception as e:
            return jsonify({"error": f"Request failed: {str(e)}"}), 502

        choice = result["choices"][0]
        assistant_msg = choice["message"]

        if "tool_calls" in assistant_msg and assistant_msg["tool_calls"]:
            tool_executions = []

            messages.append(assistant_msg)

            for tc in assistant_msg["tool_calls"]:
                tool_id = tc["id"]
                func_name = tc["function"]["name"]
                func_args = json.loads(tc["function"]["arguments"])

                if func_name == "generate_image":
                    image_tool_called = True

                exec_result = execute_tool(func_name, func_args)

                tool_exec_entry = {
                    "tool_call_id": tool_id,
                    "name": func_name,
                    "arguments": func_args,
                    "result": exec_result["result"],
                    "success": exec_result["success"]
                }
                if "filename" in exec_result:
                    tool_exec_entry["filename"] = exec_result["filename"]
                if "data_url" in exec_result:
                    tool_exec_entry["data_url"] = exec_result["data_url"]

                tool_executions.append(tool_exec_entry)

                messages.append({
                    "role": "tool",
                    "tool_call_id": tool_id,
                    "content": exec_result["result"]
                })

            pair = {
                "request": {
                    "model": "deepseek-chat",
                    "messages": [dict(m) for m in messages[:-len(assistant_msg["tool_calls"]) - 1]]
                },
                "response": {
                    "role": "assistant",
                    "content": None,
                    "tool_calls": assistant_msg["tool_calls"]
                },
                "tool_executions": tool_executions
            }
            request_response_pairs.append(pair)
            continue

        # Final text response
        reply = assistant_msg.get("content", "")

        final_response = {
            "role": "assistant",
            "content": reply
        }

        pair = {
            "request": chat_request,
            "response": final_response,
            "tool_executions": []
        }
        if request_response_pairs and request_response_pairs[-1].get("response", {}).get("content") is None:
            request_response_pairs[-1]["response"] = final_response
        else:
            request_response_pairs.append(pair)

        # Save to this conversation
        conv["messages"].append({
            "user": user_message,
            "assistant": reply,
            "timestamp": datetime.now().isoformat()
        })

        # Auto-title from first message
        if conv["title"] == "New Chat" and len(conv["messages"]) == 1:
            conv["title"] = user_message[:50]

        conv["updated"] = datetime.now().isoformat()

        # Update conv in db list
        for i, c in enumerate(db["conversations"]):
            if c["id"] == conv["id"]:
                db["conversations"][i] = conv
                break
        save_db(db)

        # Collect generated images from all tool executions
        generated_images = []
        for pair in request_response_pairs:
            for tool_exec in pair.get("tool_executions", []):
                if tool_exec.get("success") and tool_exec.get("name") == "generate_image":
                    if "filename" in tool_exec:
                        img_entry = {
                            "filename": tool_exec["filename"],
                            "prompt": tool_exec.get("arguments", {}).get("prompt", "")
                        }
                        if "data_url" in tool_exec:
                            img_entry["data_url"] = tool_exec["data_url"]
                        generated_images.append(img_entry)

        return jsonify({
            "reply": reply,
            "request_response_pairs": request_response_pairs,
            "has_tool_calls": len(request_response_pairs) > 1 or (
                len(request_response_pairs) == 1 and request_response_pairs[0].get("tool_executions", [])
            ),
            "generated_images": generated_images
        })

    return jsonify({"error": "Too many tool call iterations."}), 500


@app.route("/api/history", methods=["GET"])
def get_history():
    db = load_db()
    # Only return non-hidden conversations for UI display
    visible = [c for c in db.get("conversations", []) if not c.get("hidden")]
    return jsonify({
        "conversations": visible,
        "active": db.get("active")
    })


@app.route("/api/clear", methods=["POST"])
def clear_history():
    db = load_db()
    active_id = db.get("active")
    if active_id:
        for conv in db["conversations"]:
            if conv["id"] == active_id:
                conv["messages"] = []
                conv["updated"] = datetime.now().isoformat()
                break
        save_db(db)
    return jsonify({"status": "ok"})


@app.route("/api/conversations/<conv_id>", methods=["DELETE"])
def delete_conversation(conv_id):
    db = load_db()
    for i, conv in enumerate(db["conversations"]):
        if conv["id"] == conv_id:
            # Soft delete — hide from UI but keep data for AI memory
            conv["hidden"] = True
            conv["updated"] = datetime.now().isoformat()
            if db["active"] == conv_id:
                # Switch active to first visible conversation
                visible = [c for c in db["conversations"] if not c.get("hidden")]
                db["active"] = visible[0]["id"] if visible else None
            save_db(db)
            return jsonify({"status": "ok"})
    return jsonify({"error": "Conversation not found"}), 404


@app.route("/files/<path:filename>")
def serve_file(filename):
    from flask import send_from_directory
    return send_from_directory(str(FILES_DIR), filename)


@app.route("/upload-word", methods=["POST"])
def upload_word():
    """Upload a Word document, extract text, add to conversation context."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400
    
    ext = Path(file.filename).suffix.lower()
    if ext not in (".docx", ".doc"):
        return jsonify({"error": "Only Word documents (.docx, .doc) are supported"}), 400
    
    try:
        # Save original file
        original_name = Path(file.filename).name
        save_path = (FILES_DIR / original_name).resolve()
        if not str(save_path).startswith(str(FILES_DIR.resolve())):
            return jsonify({"error": "Invalid file path"}), 400
        save_path.parent.mkdir(parents=True, exist_ok=True)
        file.save(str(save_path))
        
        # Extract text from .docx
        if ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(save_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        else:
            # .doc format — try python-docx first, fallback to antiword or plain read
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(save_path))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n\n".join(paragraphs)
            except Exception:
                # Read as raw text (won't be perfect for .doc but gives something)
                text = save_path.read_text(encoding="utf-8", errors="replace")
        
        if not text.strip():
            return jsonify({"error": "No readable text found in the document"}), 400
        
        return jsonify({
            "status": "ok",
            "filename": original_name,
            "text": text,
            "chars": len(text),
            "preview": text[:200] + ("..." if len(text) > 200 else "")
        })
    
    except Exception as e:
        return jsonify({"error": f"Failed to process document: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)